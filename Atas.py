# Atas.py — rev7.3 (espaçamento como no modelo; cabeçalho imagem; gênero no Item 01)
from fastapi import FastAPI, Request, Form
from fastapi.responses import HTMLResponse, StreamingResponse
from starlette.responses import Response
from jinja2 import Environment, DictLoader, select_autoescape
from datetime import date, timedelta
from markupsafe import Markup
from docx.shared import Pt, Inches, Cm
from contextlib import asynccontextmanager
from docx.enum.text import WD_ALIGN_PARAGRAPH
import io, sqlite3, os

# Exportadores
from docx import Document as DocxDocument
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from xhtml2pdf import pisa

APP_TITLE = "Sistema de Atas — Comitê de Investimentos"
PARTICIPANTES = [
    "Albert Iglésia Correa dos Santos Júnior",
    "Lucas José das Neves Rodrigues",
    "Mariana Schneider Viana",
    "Shirlene Pires Mesquita",
    "Tatiana Gasparini Silva Stelzer",
]
CARGO = "Membro do Comitê de Investimentos"

# Mulheres (para “A Sra.”)
PARTICIPANTES_MULHERES = {
    "Shirlene Pires Mesquita",
    "Mariana Schneider Viana",
    "Tatiana Gasparini Silva Stelzer",
}

DB_PATH = os.getenv("DATABASE_URL", "./atas.db")
BASE_DIR = os.path.dirname(os.path.abspath(__file__))
STATIC_DIR = os.path.join(BASE_DIR, "static")

TOPICOS = [
    ("CHINA", "CHINA"),
    ("ESTADOS UNIDOS", "ESTADOS UNIDOS"),
    ("EUROPA", "EUROPA"),
    ("CENÁRIO POLÍTICO BRASILEIRO", "CENÁRIO POLÍTICO BRASILEIRO"),
    ("CENÁRIO ECONÔMICO BRASILEIRO", "CENÁRIO ECONÔMICO BRASILEIRO"),
]

# ===== Espaçamento (pt) =====
HEADER_GAP  = 12  # antes do primeiro título, abaixo do cabeçalho
TITLE_GAP   = 6   # depois de cada título
SECTION_GAP = 12  # entre blocos principais

# --------------------- Datas
def ptbr_date(d: date) -> str:
    return d.strftime("%d/%m/%Y")

def mes_ano_pt(d: date) -> str:
    nomes = {1:"janeiro",2:"fevereiro",3:"março",4:"abril",5:"maio",6:"junho",7:"julho",8:"agosto",9:"setembro",10:"outubro",11:"novembro",12:"dezembro"}
    return f"{nomes[d.month]} de {d.year}"

def menos_dois_meses(d: date) -> date:
    m = d.month - 2; y = d.year
    if m <= 0: m += 12; y -= 1
    dia = min(d.day, 28)
    return date(y, m, dia)

def primeira_quinta(ano: int, mes: int) -> date:
    d = date(ano, mes, 1)
    offset = (3 - d.weekday()) % 7  # quinta=3
    return d + timedelta(days=offset)

# --------------------- DB
def db_connect():
    conn = sqlite3.connect(DB_PATH)
    conn.row_factory = sqlite3.Row
    return conn

def ensure_schema():
    conn = db_connect(); cur = conn.cursor()
    cur.execute("""
        CREATE TABLE IF NOT EXISTS meeting (
            id INTEGER PRIMARY KEY,
            numero INTEGER NOT NULL,
            ano INTEGER NOT NULL,
            data TEXT NOT NULL,
            hora TEXT NOT NULL,
            local TEXT NOT NULL
        );
    """)
    cur.execute("""
        CREATE TABLE IF NOT EXISTS scenario (
            id INTEGER PRIMARY KEY,
            participant TEXT NOT NULL,
            text TEXT NOT NULL DEFAULT ''
        );
    """)
    # migrações
    cur.execute("PRAGMA table_info('scenario')")
    cols = [r[1] for r in cur.fetchall()]
    if "topic" not in cols:
        cur.execute("ALTER TABLE scenario ADD COLUMN topic TEXT NOT NULL DEFAULT ''")
    cur.execute("PRAGMA table_info('meeting')")
    mcols = [r[1] for r in cur.fetchall()]
    if "lavrador" not in mcols:
        cur.execute("ALTER TABLE meeting ADD COLUMN lavrador TEXT NOT NULL DEFAULT ''")

    cur.execute("SELECT COUNT(1) FROM meeting")
    if cur.fetchone()[0] == 0:
        hoje = date.today(); pq = primeira_quinta(hoje.year, hoje.month)
        cur.execute(
            "INSERT INTO meeting (numero, ano, data, hora, local, lavrador) VALUES (?, ?, ?, ?, ?, ?)",
            (4, hoje.year, pq.isoformat(), "14:00", "Sala nº 408 do 4º andar do IPAJM", ""),
        )
    cur.execute("SELECT COUNT(1) FROM scenario")
    if cur.fetchone()[0] == 0:
        for p in PARTICIPANTES:
            cur.execute("INSERT INTO scenario (participant, text, topic) VALUES (?, '', '')", (p,))
    conn.commit(); conn.close()

# --------------------- HTML TEMPLATES (UI + Prévia/PDF)
TEMPLATES = {
    "base.html": r"""<!doctype html><html lang="pt-br"><head>
  <meta charset="utf-8"/><meta name="viewport" content="width=device-width, initial-scale=1"/>
  <title>{{ title }}</title>
  <script src="https://unpkg.com/htmx.org@1.9.12"></script>
  <script src="https://cdn.tailwindcss.com"></script>
  <style>
    body{font-family: Inter, system-ui, -apple-system, Segoe UI, Roboto, Helvetica, Arial, sans-serif; font-size:16px; color:#111827}
    .h-bold{font-weight:700}
    .prose.ata-preview{font-family: Calibri, "Segoe UI", Arial, sans-serif !important; font-size:11pt !important; color:#000 !important; text-align:justify; line-height:1.35; white-space:pre-wrap;}
    .prose.ata-preview strong{font-weight:700}
  </style>
</head><body class="bg-gray-50 text-gray-900">
  <header class="sticky top-0 z-10 bg-white shadow-sm">
    <div class="max-w-5xl mx-auto px-3 py-3 flex items-center justify-between">
      <h1 class="text-lg h-bold">Sistema de Atas — Comitê de Investimentos</h1>
      <div class="flex gap-2">
        <a href="/export/docx" class="px-3 py-2 rounded bg-indigo-600 text-white hover:bg-indigo-700">Exportar DOCX</a>
        
      </div>
    </div>
  </header>
  <main class="max-w-5xl mx-auto px-3 py-5 grid md:grid-cols-2 gap-6">{% block content %}{% endblock %}</main>
  <footer class="max-w-5xl mx-auto px-3 pb-8 text-xs text-gray-700">Prévia/PDF: Calibri 11 e cabeçalho idêntico. DOCX: gerado programaticamente com mesmo layout.</footer>
</body></html>
""",
    "index.html": r"""{% extends 'base.html' %}{% block content %}
  {% include 'partials/session.html' %}

  <section class="bg-white rounded-xl shadow p-4">
    <h2 class="text-base h-bold mb-3">Item 01 — Cenário (por analista)</h2>
    <div id="item1-form">{% include 'partials/item1_form.html' with context %}</div>
    <div id="status-list">{% include 'partials/status.html' %}</div>
  </section>

  <section class="bg-white rounded-xl shadow p-4">
    <h2 class="text-base h-bold mb-3">Itens 02 e 04 — Texto livre</h2>
    <form hx-post="/preview/update" hx-target="this" hx-swap="none" class="grid gap-3">
      <div><label class="text-sm h-bold">Item 02 — Movimentações e Aplicações financeiras</label>
        <textarea name="item2" class="min-h-[100px] w-full px-3 py-2 rounded border">{{ item2 }}</textarea></div>
      <div><label class="text-sm h-bold">Item 04 — Assuntos Gerais</label>
        <textarea name="assuntos" class="min-h-[120px] w-full px-3 py-2 rounded border">{{ assuntos }}</textarea></div>
      <button class="justify-self-start px-4 py-2 rounded shadow bg-slate-700 text-white hover:bg-slate-800">Salvar Itens 02 e 04</button>
    </form>
  </section>

  <section class="bg-white rounded-xl shadow p-4">
    <h2 class="text-base h-bold mb-3">Item 03 — Parâmetros</h2>
    <form hx-post="/resumo/update" hx-target="this" hx-swap="none" class="grid gap-3">
      <div class="grid md:grid-cols-4 gap-3">
        <div class="md:col-span-2"><label class="text-sm">Rentabilidade do Fundo ( % ) no mês D-2</label>
          <input type="text" name="rentab" placeholder="ex: 1,23%" class="w-full px-3 py-2 rounded border"/></div>
        <div><label class="text-sm">Diferença vs. meta (p.p.)</label>
          <input type="text" name="difpp" placeholder="ex: 0,15" class="w-full px-3 py-2 rounded border"/></div>
        <div><label class="text-sm">Posição</label>
          <select name="posicao" class="w-full px-3 py-2 rounded border"><option value="abaixo">abaixo</option><option value="acima">acima</option></select></div>
      </div>
      <div class="grid md:grid-cols-4 gap-3">
        <div><label class="text-sm">Risco assumido ( % ) no mês D-2</label>
          <input type="text" name="risco" placeholder="ex: 7,5%" class="w-full px-3 py-2 rounded border"/></div>
      </div>
      <button class="justify-self-start px-4 py-2 rounded shadow bg-slate-700 text-white hover:bg-slate-800">Salvar Parâmetros do Item 03</button>
    </form>
    <p class="text-xs text-gray-600 mt-2">* D-2 = dois meses antes da data da reunião.</p>
  </section>

  <section class="bg-white rounded-xl shadow p-4 md:col-span-2">
    <div class="flex items-center justify-between"><h2 class="text-base h-bold">Prévia da Ata</h2>
      <button class="px-4 py-2 rounded shadow bg-purple-600 text-white hover:bg-purple-700"
              hx-get="/preview/modal" hx-target="body" hx-swap="beforeend">Ver prévia da ata</button></div>
    <p class="text-xs text-gray-600 mt-2">A prévia abre em uma janela com rolagem.</p>
  </section>
{% endblock %}
""",
    "partials/item1_form.html": r"""
<form id="item1-form" hx-post="/scenario/save" hx-target="#status-list" hx-swap="outerHTML" class="grid gap-3">
  <div class="grid md:grid-cols-3 gap-3">
    <div><label class="text-sm">Quem está preenchendo?</label>
      <select name="participant" class="w-full px-3 py-2 rounded border">
        {% for p in participantes %}<option value="{{p}}" {{ 'selected' if p==participant else '' }}>{{p}}</option>{% endfor %}
      </select>
    </div>
    <div><label class="text-sm">Tema</label>
      <select name="topic" class="w-full px-3 py-2 rounded border">
        {% for val,label in topicos %}<option value="{{val}}" {{ 'selected' if val==topic else '' }}>{{label}}</option>{% endfor %}
      </select>
    </div>
    <div class="flex items-end"><button class="px-4 py-2 rounded shadow bg-blue-600 text-white hover:bg-blue-700">Salvar Parte</button></div>
  </div>
  <textarea name="text" placeholder="Cole o texto do seu cenário…" class="min-h-[180px] w-full px-3 py-2 rounded border">{{ text }}</textarea>
</form>
""",
    "partials/session.html": r"""
<section id="session-block" class="bg-white rounded-xl shadow p-4 md:col-span-2">
  <h2 class="text-base h-bold mb-3">Configurações da Ata</h2>
  <form hx-post="/meeting/update" hx-target="#session-block" hx-swap="outerHTML" class="grid md:grid-cols-6 gap-3">
    <div class="col-span-2">
      <label class="text-sm">Nº da Ata</label>
      <div class="flex gap-2">
        <input type="number" name="numero" value="{{ meeting.numero }}" class="w-24 px-3 py-2 rounded border"/>
        <input type="number" name="ano" value="{{ meeting.ano }}" class="w-28 px-3 py-2 rounded border"/>
      </div>
      <p class="text-xs mt-1">Formato exibido: {{ '%03d' % meeting.numero }}/{{ meeting.ano }}</p>
    </div>
    <div class="col-span-2"><label class="text-sm">Data</label>
      <input type="date" name="data" value="{{ meeting.data }}" class="w-full px-3 py-2 rounded border"/></div>
    <div><label class="text-sm">Hora</label>
      <input type="time" name="hora" value="{{ meeting.hora }}" class="w-full px-3 py-2 rounded border"/></div>
    <div class="col-span-3"><label class="text-sm">Local</label>
      <input type="text" name="local" value="{{ meeting.local }}" class="w-full px-3 py-2 rounded border"/></div>
    <div class="col-span-3"><label class="text-sm">Responsável pela ata</label>
      <select name="lavrador" class="w-full px-3 py-2 rounded border">
        <option value="">(selecionar)</option>
        {% for p in participantes %}<option value="{{p}}" {{ 'selected' if meeting.lavrador==p else '' }}>{{p}}</option>{% endfor %}
      </select>
    </div>
    <div class="col-span-3 flex items-end"><button class="px-4 py-2 rounded shadow bg-blue-600 text-white hover:bg-blue-700">Salvar</button></div>
  </form>
</section>
""",
    "partials/status.html": r"""
<div id="status-list">
  <h3 class="font-medium mb-2">Status dos participantes</h3>
  <ul class="grid grid-cols-1 md:grid-cols-2 gap-2">
    {% for p in participantes %}
      {% set ok = scenarios.get(p, {}).get('text','').strip()|length > 0 %}
      <li class="flex items-center justify-between px-3 py-2 rounded border {{ 'bg-emerald-50 border-emerald-200' if ok else 'bg-gray-50' }}">
        <span class="text-sm">{{ p }}</span>
        <div class="flex items-center gap-2">
          <span class="text-xs px-2 py-1 rounded-full {{ 'bg-emerald-200 text-emerald-900' if ok else 'bg-gray-200 text-gray-700' }}">{{ 'Preenchido' if ok else 'Pendente' }}</span>
          <button class="text-xs px-2 py-1 rounded bg-slate-200 hover:bg-slate-300"
                  hx-get="/scenario/form?participant={{p|urlencode}}"
                  hx-target="#item1-form" hx-swap="outerHTML">Editar</button>
        </div>
      </li>
    {% endfor %}
  </ul>
</div>
""",
    "partials/preview_modal.html": r"""
<div id="preview-modal" class="fixed inset-0 z-50">
  <div class="absolute inset-0 bg-black/50" onclick="document.getElementById('preview-modal')?.remove()"></div>
  <div class="absolute inset-0 flex items-center justify-center p-4">
    <div class="relative bg-white rounded-2xl shadow-xl w-[95vw] max-w-5xl max-h-[85vh] overflow-hidden border flex flex-col">
      <div class="flex items-center justify-between px-4 py-3 border-b bg-gray-50 shrink-0">
        <h3 class="text-base font-semibold">Prévia da Ata</h3>
        <button class="px-3 py-1.5 rounded bg-gray-200 hover:bg-gray-300" onclick="document.getElementById('preview-modal')?.remove()">Fechar</button>
      </div>
      <div class="grow overflow-y-auto p-5 overscroll-contain">
        <div class="prose ata-preview max-w-none">{{ ata_html | safe }}</div>
      </div>
    </div>
  </div>
</div>
<script>(function(){const onEsc=(e)=>{if(e.key==='Escape'){document.getElementById('preview-modal')?.remove();document.removeEventListener('keydown',onEsc);}};document.addEventListener('keydown',onEsc);})();</script>
""",
}

env = Environment(loader=DictLoader(TEMPLATES), autoescape=select_autoescape(["html","xml"]))
def render(name: str, **ctx) -> HTMLResponse:
    return HTMLResponse(env.get_template(name).render(**ctx))

def nl2br(value: str) -> Markup:
    if not value: return Markup("")
    return Markup(value.replace("\n", "<br/>"))
env.filters["nl2br"] = nl2br

# --------------------- Estado
ITEM2_DEFAULT = "Não houve realocações de recursos desde a última reunião até a presente data."
ASSUNTOS_DEFAULT = "– Assuntos gerais discutidos e/ou Eventos:"
STATE = {"item2": ITEM2_DEFAULT, "assuntos": ASSUNTOS_DEFAULT, "resumo": {"rentab": "", "difpp": "", "posicao": "abaixo", "risco": ""}}

# --------------------- App
app = FastAPI(title=APP_TITLE)

@asynccontextmanager
async def lifespan(app: FastAPI):
    os.makedirs(STATIC_DIR, exist_ok=True)
    ensure_schema()
    yield

app = FastAPI(title=APP_TITLE, lifespan=lifespan)

# --------------------- Helpers (HTML/PDF)
def numero_sessao_fmt(numero: int, ano: int) -> str:
    return f"{numero:03d}/{ano}"

def _prefixo_genero(nome: str) -> str:
    return "A Sra." if nome in PARTICIPANTES_MULHERES else "O Sr."

def item1_single_paragraph(meeting: dict, scenarios: dict) -> str:
    d = date.fromisoformat(meeting["data"])

    ORDINAIS = {
        1: "primeiro",
        2: "segundo",
        3: "terceiro",
        4: "quarto",
        5: "quinto",
        6: "sexto",
        7: "sétimo",
        8: "oitavo",
        9: "nono",
        10: "décimo",
        11: "décimo primeiro",
        12: "décimo segundo",
        13: "décimo terceiro",
        14: "décimo quarto",
        15: "décimo quinto",
        16: "décimo sexto",
        17: "décimo sétimo",
        18: "décimo oitavo",
        19: "décimo nono",
        20: "vigésimo",
        21: "vigésimo primeiro",
        22: "vigésimo segundo",
        23: "vigésimo terceiro",
        24: "vigésimo quarto",
        25: "vigésimo quinto",
        26: "vigésimo sexto",
        27: "vigésimo sétimo",
        28: "vigésimo oitavo",
        29: "vigésimo nono",
        30: "trigésimo",
        31: "trigésimo primeiro"
    }

    intro = (
        f"No {ORDINAIS.get(d.day, str(d.day))} dia do mês de {mes_ano_pt(d).split(' de ')[0]} "
        f"do ano de {d.year}, às {meeting['hora']} horas, na {meeting['local']}, "
        f"ocorreu a {meeting['numero']}ª Reunião Ordinária dos Membros do Comitê de Investimentos. "
    )

    partes = [intro]
    for nome in PARTICIPANTES:
        row = scenarios.get(nome, {"text": "", "topic": ""})
        texto = (row.get("text") or "").strip()
        tema = (row.get("topic") or "").strip()
        if texto:
            partes.append(
                f"<strong>{_prefixo_genero(nome)} {nome}</strong> falando sobre {tema}, {texto} "
            )

    return "<p>" + "".join(partes).strip() + "</p>"


def item3_html(meeting: dict) -> str:
    d = date.fromisoformat(meeting["data"]); d2 = menos_dois_meses(d); mes_ano = mes_ano_pt(d2)
    p0 = ("O Comitê de Investimentos, buscando transmitir maior transparência em relação às análises dos investimentos do Instituto e, em consequência, "
          "aderindo às normas do Pró-Gestão, elabora o “Relatório de Análise de Investimentos IPAJM”. "
          "Este relatório já foi encaminhado à SCO – Subgerência de Contabilidade e Orçamento, para posterior envio para análise do Conselho Fiscal do IPAJM. "
          f"Segue abaixo um resumo relativo aos itens abordados no Relatório supracitado de {mes_ano}:")
    r = STATE["resumo"]; rentab=r["rentab"]; difpp=r["difpp"]; pos=r["posicao"]; risco=r["risco"]
    p1 = f"1) Acompanhamento da rentabilidade -  A rentabilidade consolidada dos investimentos do Fundo Previdenciário em {mes_ano} foi de {rentab}, ficando {difpp} p.p. {pos} da meta atuarial."
    p2 = f"2) Avaliação de risco da carteira - O grau de variação nas rentabilidades está coerente com o grau de risco assumido, em {risco}."
    p3 = f"3) Execução da Política de Investimentos – As movimentações financeiras realizadas no mês de {mes_ano} estão de acordo com as deliberações estabelecidas com a Diretoria de Investimentos e com a legislação vigente."
    p4 = f"4) Aderência a Política de Investimentos - Os recursos investidos, abrangendo a carteira consolidada, que representa o patrimônio total do RPPS sob gestão, estão aderentes à Política de Investimentos de {meeting['ano']}, respeitando o estabelecido na legislação em vigor e dentro dos percentuais definidos.  Considerando que as taxas ainda são negociadas acima da meta atuarial, seguimos com a estratégia de alcançar o alvo definido de 60% de alocação em Títulos Públicos."
    return "\n".join([f"<p><strong>{t}</strong></p>" if i==0 else f"<p>{t}</p>" for i,t in enumerate([p0,p1,p2,p3,p4])])

def ata_html_full(meeting: dict, scenarios: dict) -> str:
    d = date.fromisoformat(meeting["data"]); numero_fmt = numero_sessao_fmt(meeting["numero"], meeting["ano"])
    presencas_html = "<br/>".join([f"<strong>{p}</strong> - {CARGO};" for p in PARTICIPANTES])
    blocos = [
        f"<p><strong>Sessão Ordinária nº {numero_fmt}</strong></p>",
        f"<p><strong>Data:</strong> {ptbr_date(d)}.<br/><strong>Hora:</strong> {meeting['hora']}h.<br/><strong>Local:</strong> {meeting['local']}.</p>",
        "<p><strong>Presenças:</strong></p>",
        f"<p>{presencas_html}</p>",
        "<p><strong>Ordem do Dia:</strong></p>",
        "<p>1. <strong>Cenário Político e Econômico Interno</strong> e <strong>Cenário Econômico Externo (EUA, Europa e China)</strong>;<br/>"
        "2. <strong>Movimentações e Aplicações financeiras</strong>;<br/>"
        "3. <strong>Acompanhamento dos Recursos Investidos</strong>;<br/>"
        "4. <strong>Assuntos Gerais</strong>.</p>",
        "<p><strong>Item 01 – Cenário Político e Econômico Interno e Cenário Econômico Externo (EUA, Europa e China):</strong></p>",
        item1_single_paragraph(meeting, scenarios),
        "<p><strong>Item 02 – Movimentações e Aplicações financeiras</strong></p>",
        f"<p>{STATE['item2']}</p>",
        "<p><strong>Item 03 – Acompanhamento dos Recursos Investidos:</strong></p>",
        item3_html(meeting),
        "<p><strong>Item 04 – Assuntos Gerais</strong></p>",
        f"<p>{STATE['assuntos']}</p>",
        (lambda lav: f"<p>Nada mais havendo a tratar, foi encerrada a reunião e eu, "
                     f"{f'<strong>{lav}</strong>' if lav else '___________________________________'}, "
                     "lavrei a presente Ata, assinada pelos membros presentes do Comitê de Investimentos.</p>")((meeting.get('lavrador') or '').strip()),
        f"<p><strong>{PARTICIPANTES[0]}</strong>&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;<strong>{PARTICIPANTES[1]}</strong><br/>{CARGO}&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;{CARGO}</p>",
        f"<p><strong>{PARTICIPANTES[2]}</strong>&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;<strong>{PARTICIPANTES[3]}</strong><br/>{CARGO}&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;{CARGO}</p>",
        f"<p><strong>{PARTICIPANTES[4]}</strong><br/>{CARGO}</p>",
    ]
    # Cabeçalho HTML/PDF
    brasao = os.path.join(STATIC_DIR, "brasao.png")
    simbolo = os.path.join(STATIC_DIR, "simbolo.png")
    header_html = f"""
    <div style="width:100%; display:flex; align-items:center; justify-content:space-between; margin-bottom:8px;">
      <img src="{brasao}" style="height:42px"/>
      <div style="text-align:center; font-family: Calibri; font-size:11pt; font-weight:700; color:#000;">
        GOVERNO DO ESTADO DO ESPÍRITO SANTO<br/>INSTITUTO DE PREVIDÊNCIA DOS<br/>SERVIDORES DO ESTADO DO ESPÍRITO SANTO
      </div>
      <img src="{simbolo}" style="height:42px"/>
    </div>
    <div style="border-bottom:1px solid #000; margin:2px 0 8px 0; text-align:center;">IPAJM</div>
    """
    return header_html + "\n".join(blocos)

# --------------------- ROTAS (UI)
@app.get("/", response_class=HTMLResponse)
async def index(request: Request):
    conn = db_connect(); cur = conn.cursor()
    cur.execute("SELECT * FROM meeting LIMIT 1"); meeting = dict(cur.fetchone())
    cur.execute("SELECT participant, text, topic FROM scenario")
    scenarios = {r[0]: {"text": r[1], "topic": r[2]} for r in cur.fetchall()}
    conn.close()
    ata_html = ata_html_full(meeting, scenarios)
    return render("index.html", title=APP_TITLE, participantes=PARTICIPANTES, topicos=TOPICOS,
                  meeting=meeting, scenarios=scenarios, item2=STATE["item2"], assuntos=STATE["assuntos"], ata_html=ata_html)

@app.post("/meeting/update")
async def meeting_update(numero: int = Form(...), ano: int = Form(...), data: str = Form(...),
                         hora: str = Form(...), local: str = Form(...), lavrador: str = Form("")):
    conn = db_connect(); cur = conn.cursor()
    cur.execute("SELECT id FROM meeting LIMIT 1"); mid = cur.fetchone()[0]
    cur.execute("UPDATE meeting SET numero=?, ano=?, data=?, hora=?, local=?, lavrador=? WHERE id=?",
                (numero, ano, data, hora, local, lavrador, mid))
    conn.commit()
    cur.execute("SELECT * FROM meeting WHERE id=?", (mid,)); meeting = dict(cur.fetchone()); conn.close()
    html = env.get_template("partials/session.html").render(meeting=meeting, participantes=PARTICIPANTES)
    return HTMLResponse(html)

@app.get("/scenario/form")
async def scenario_form(participant: str):
    conn = db_connect(); cur = conn.cursor()
    cur.execute("SELECT participant, text, topic FROM scenario WHERE participant=?", (participant,))
    row = cur.fetchone(); conn.close()
    text = row["text"] if row else ""; topic = row["topic"] if row else ""
    html = env.get_template("partials/item1_form.html").render(participantes=PARTICIPANTES, topicos=TOPICOS,
                                                               participant=participant, topic=topic, text=text)
    return HTMLResponse(html)

@app.post("/scenario/save")
async def scenario_save(participant: str = Form(...), topic: str = Form(...), text: str = Form(...)):
    conn = db_connect(); cur = conn.cursor()
    cur.execute("UPDATE scenario SET text=?, topic=? WHERE participant=?", (text, topic, participant))
    conn.commit()
    cur.execute("SELECT participant, text, topic FROM scenario")
    scenarios = {r[0]: {"text": r[1], "topic": r[2]} for r in cur.fetchall()}
    conn.close()
    status_html = env.get_template("partials/status.html").render(participantes=PARTICIPANTES, scenarios=scenarios)
    return HTMLResponse(f'<div id="status-list">{status_html}</div>')

@app.post("/preview/update")
async def preview_update(item2: str = Form(...), assuntos: str = Form(...)):
    STATE["item2"] = item2; STATE["assuntos"] = assuntos
    return Response(status_code=204)

@app.post("/resumo/update")
async def resumo_update(rentab: str = Form(""), difpp: str = Form(""), posicao: str = Form("abaixo"), risco: str = Form("")):
    STATE["resumo"].update({"rentab": rentab, "difpp": difpp, "posicao": posicao, "risco": risco})
    return Response(status_code=204)

@app.get("/preview/modal")
async def preview_modal():
    conn = db_connect(); cur = conn.cursor()
    cur.execute("SELECT * FROM meeting LIMIT 1"); meeting = dict(cur.fetchone())
    cur.execute("SELECT participant, text, topic FROM scenario")
    scenarios = {r[0]: {"text": r[1], "topic": r[2]} for r in cur.fetchall()}
    conn.close()
    ata_html = ata_html_full(meeting, scenarios)
    html = env.get_template("partials/preview_modal.html").render(ata_html=ata_html)
    return HTMLResponse(html)

# --------------------- Exportações
@app.get("/export/pdf")
async def export_pdf():
    conn = db_connect(); cur = conn.cursor()
    cur.execute("SELECT * FROM meeting LIMIT 1"); meeting = dict(cur.fetchone())
    cur.execute("SELECT participant, text, topic FROM scenario")
    scenarios = {r[0]: {"text": r[1], "topic": r[2]} for r in cur.fetchall()}
    conn.close()
    html_body = ata_html_full(meeting, scenarios)
    html = f"""<html><head><meta charset='utf-8'>
      <style>body {{ font-family: Calibri, Arial; font-size: 11pt; color:#000; }}
        .content {{ white-space: normal; text-align: justify; }}</style>
    </head><body><div class="content">{html_body}</div></body></html>"""
    pdf_buf = io.BytesIO(); pisa.CreatePDF(io.StringIO(html), dest=pdf_buf)
    pdf_buf.seek(0)
    filename = f"Ata_{meeting['numero']:03d}-{meeting['ano']}.pdf"
    return StreamingResponse(pdf_buf, media_type="application/pdf",
                             headers={"Content-Disposition": f"attachment; filename={filename}"})

# ---------- DOCX programático ----------
def _set_cell_borders(cell, bottom=True):
    tc = cell._tc; tcPr = tc.get_or_add_tcPr(); tcBorders = OxmlElement("w:tcBorders")
    if bottom:
        bottom_el = OxmlElement("w:bottom"); bottom_el.set(qn("w:val"), "single"); bottom_el.set(qn("w:sz"), "8"); bottom_el.set(qn("w:color"), "000000")
        tcBorders.append(bottom_el)
    tcPr.append(tcBorders)

def _add_header_image(doc: DocxDocument, path=os.path.join(STATIC_DIR, "cabecalho.png"), width_inches=6.5):
    if not os.path.exists(path): return
    header = doc.sections[0].header
    p = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(path, width=Inches(width_inches))
    p.paragraph_format.space_after = Pt(0)

def _add_paragraph(doc, text, bold=False, align=WD_ALIGN_PARAGRAPH.JUSTIFY, space_after_pt=0):
    p = doc.add_paragraph()
    run = p.add_run(text); run.font.name = "Calibri"; run.font.size = Pt(11); run.bold = bold
    p.alignment = align
    if space_after_pt: p.paragraph_format.space_after = Pt(space_after_pt)
    return p

def _add_label_value(doc, label, value, space_after_pt=0):
    p = doc.add_paragraph()
    r1 = p.add_run(f"{label}: "); r1.bold = True; r1.font.name = "Calibri"; r1.font.size = Pt(11)
    r2 = p.add_run(value); r2.font.name = "Calibri"; r2.font.size = Pt(11)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    if space_after_pt: p.paragraph_format.space_after = Pt(space_after_pt)
    return p

def _add_title(doc, text, space_after_pt=TITLE_GAP):
    return _add_paragraph(doc, text, bold=True, align=WD_ALIGN_PARAGRAPH.LEFT, space_after_pt=space_after_pt)

def _prefixo_genero_docx(nome: str) -> str:
    return "A Sra." if nome in PARTICIPANTES_MULHERES else "O Sr."

def _item01_single_paragraph_docx(doc, meeting, scenarios, space_after_pt=SECTION_GAP):
    d = date.fromisoformat(meeting["data"])
    ORDINAIS = {
        1: "primeiro",
        2: "segundo",
        3: "terceiro",
        4: "quarto",
        5: "quinto",
        6: "sexto",
        7: "sétimo",
        8: "oitavo",
        9: "nono",
        10: "décimo",
        11: "décimo primeiro",
        12: "décimo segundo",
        13: "décimo terceiro",
        14: "décimo quarto",
        15: "décimo quinto",
        16: "décimo sexto",
        17: "décimo sétimo",
        18: "décimo oitavo",
        19: "décimo nono",
        20: "vigésimo",
        21: "vigésimo primeiro",
        22: "vigésimo segundo",
        23: "vigésimo terceiro",
        24: "vigésimo quarto",
        25: "vigésimo quinto",
        26: "vigésimo sexto",
        27: "vigésimo sétimo",
        28: "vigésimo oitavo",
        29: "vigésimo nono",
        30: "trigésimo",
        31: "trigésimo primeiro"
    }
    intro = (
        f"No {ORDINAIS.get(d.day, str(d.day))} dia do mês de {mes_ano_pt(d).split(' de ')[0]} "
        f"às {meeting['hora']} horas, na {meeting['local']}, ocorreu a {meeting['numero']}ª "
        f"Reunião Ordinária dos Membros do Comitê de Investimentos. "
    )
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    r = p.add_run(intro); r.font.name = "Calibri"; r.font.size = Pt(11)
    for nome in PARTICIPANTES:
        row = scenarios.get(nome, {"text": "", "topic": ""})
        texto = (row.get("text") or "").strip(); tema = (row.get("topic") or "").strip()
        if texto:
            pref = _prefixo_genero_docx(nome)
            rb = p.add_run(f"{pref} {nome}"); rb.bold = True; rb.font.name = "Calibri"; rb.font.size = Pt(11)
            p.add_run(f" falando sobre {tema}, {texto} ").font.size = Pt(11)
    p.paragraph_format.space_after = Pt(space_after_pt)

def _assinaturas_table(doc: DocxDocument):
    t = doc.add_table(rows=3, cols=2); t.autofit = True
    # linha 1
    cells = t.rows[0].cells
    for idx, nome in enumerate([PARTICIPANTES[0], PARTICIPANTES[1]]):
        p1 = cells[idx].paragraphs[0]; p1.alignment = WD_ALIGN_PARAGRAPH.LEFT
        r = p1.add_run(nome); r.bold = True; r.font.name = "Calibri"; r.font.size = Pt(11)
        p2 = cells[idx].add_paragraph(); r2 = p2.add_run(CARGO); r2.font.name = "Calibri"; r2.font.size = Pt(11)
    # linha 2
    cells = t.rows[1].cells
    for idx, nome in enumerate([PARTICIPANTES[2], PARTICIPANTES[3]]):
        p1 = cells[idx].paragraphs[0]; p1.alignment = WD_ALIGN_PARAGRAPH.LEFT
        r = p1.add_run(nome); r.bold = True; r.font.name = "Calibri"; r.font.size = Pt(11)
        p2 = cells[idx].add_paragraph(); r2 = p2.add_run(CARGO); r2.font.name = "Calibri"; r2.font.size = Pt(11)
    # linha 3 mesclada
    c = t.rows[2].cells; c[0].merge(c[1])
    p1 = c[0].paragraphs[0]; p1.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r = p1.add_run(PARTICIPANTES[4]); r.bold = True; r.font.name = "Calibri"; r.font.size = Pt(11)
    p2 = c[0].add_paragraph(); rr = p2.add_run(CARGO); rr.font.name = "Calibri"; rr.font.size = Pt(11)

@app.get("/export/docx")
async def export_docx():
    conn = db_connect(); cur = conn.cursor()
    cur.execute("SELECT * FROM meeting LIMIT 1"); meeting = dict(cur.fetchone())
    cur.execute("SELECT participant, text, topic FROM scenario")
    scenarios = {r[0]: {"text": r[1], "topic": r[2]} for r in cur.fetchall()}
    conn.close()

    d = date.fromisoformat(meeting["data"]); numero_fmt = numero_sessao_fmt(meeting["numero"], meeting["ano"])

    doc = DocxDocument()
    style = doc.styles["Normal"]; style.font.name = "Calibri"; style.font.size = Pt(11)

    # Cabeçalho (imagem) + espaço antes do primeiro título
    _add_header_image(doc, os.path.join(STATIC_DIR, "cabecalho.png"))

    # Sessão
    p_titulo = _add_title(doc, f"Sessão Ordinária nº {numero_fmt}", space_after_pt=TITLE_GAP)
    p_titulo.paragraph_format.space_before = Pt(HEADER_GAP)

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)  # controla a altura do espaço

    # Data / Hora / Local
    _add_label_value(doc, "Data", ptbr_date(d), space_after_pt=0)
    _add_label_value(doc, "Hora", f"{meeting['hora']}h", space_after_pt=0)
    _add_label_value(doc, "Local", meeting["local"], space_after_pt=SECTION_GAP)

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(10)  # controla a altura do espaço

    # Presenças
    _add_title(doc, "Presenças:", space_after_pt=TITLE_GAP)
    last_p = None
    for nome in PARTICIPANTES:
        p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run_nome = p.add_run(nome); run_nome.font.name = "Calibri"; run_nome.font.size = Pt(11)
        run_cargo = p.add_run(f" - {CARGO};"); run_cargo.font.name = "Calibri"; run_cargo.font.size = Pt(11)
        last_p = p
    if last_p: last_p.paragraph_format.space_after = Pt(SECTION_GAP)

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(10)  # controla a altura do espaço

    # Ordem do Dia
    _add_title(doc, "Ordem do Dia:", space_after_pt=TITLE_GAP)
    _add_paragraph(doc, "1. Cenário Político e Econômico Interno e Cenário Econômico Externo (EUA, Europa e China);", False, WD_ALIGN_PARAGRAPH.LEFT, space_after_pt=0)
    _add_paragraph(doc, "2. Movimentações e Aplicações financeiras;", False, WD_ALIGN_PARAGRAPH.LEFT, space_after_pt=0)
    _add_paragraph(doc, "3. Acompanhamento dos Recursos Investidos;", False, WD_ALIGN_PARAGRAPH.LEFT, space_after_pt=0)
    _add_paragraph(doc, "4. Assuntos Gerais.", False, WD_ALIGN_PARAGRAPH.LEFT, space_after_pt=SECTION_GAP)

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(10)  # controla a altura do espaço

    # Item 01
    _add_title(doc, "Item 01 – Cenário Político e Econômico Interno e Cenário Econômico Externo (EUA, Europa e China):", space_after_pt=TITLE_GAP)
    _item01_single_paragraph_docx(doc, meeting, scenarios, space_after_pt=SECTION_GAP)

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(10)  # controla a altura do espaço

    # Item 02
    _add_title(doc, "Item 02 – Movimentações e Aplicações financeiras", space_after_pt=TITLE_GAP)
    _add_paragraph(doc, STATE["item2"], False, WD_ALIGN_PARAGRAPH.JUSTIFY, space_after_pt=SECTION_GAP)

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(10)  # controla a altura do espaço

    # Item 03
    _add_title(doc, "Item 03 – Acompanhamento dos Recursos Investidos:", space_after_pt=TITLE_GAP)
    d2 = menos_dois_meses(d); mes_ano = mes_ano_pt(d2)
    p0 = ("O Comitê de Investimentos, buscando transmitir maior transparência em relação às análises dos investimentos do Instituto e, em consequência, "
          "aderindo às normas do Pró-Gestão, elabora o “Relatório de Análise de Investimentos IPAJM”. Este relatório já foi encaminhado à SCO – Subgerência de Contabilidade e Orçamento, "
          "para posterior envio para análise do Conselho Fiscal do IPAJM. "
          f"Segue abaixo um resumo relativo aos itens abordados no Relatório supracitado de {mes_ano}:")
    _add_paragraph(doc, p0, False, WD_ALIGN_PARAGRAPH.JUSTIFY, space_after_pt=0)
    r = STATE["resumo"]; rentab=r["rentab"]; difpp=r["difpp"]; pos=r["posicao"]; risco=r["risco"]
    _add_paragraph(doc, f"1) Acompanhamento da rentabilidade -  A rentabilidade consolidada dos investimentos do Fundo Previdenciário em {mes_ano} foi de {rentab}, ficando {difpp} p.p. {pos} da meta atuarial.", False, space_after_pt=0)
    _add_paragraph(doc, f"2) Avaliação de risco da carteira - O grau de variação nas rentabilidades está coerente com o grau de risco assumido, em {risco}.", False, space_after_pt=0)
    _add_paragraph(doc, f"3) Execução da Política de Investimentos – As movimentações financeiras realizadas no mês de {mes_ano} estão de acordo com as deliberações estabelecidas com a Diretoria de Investimentos e com a legislação vigente.", False, space_after_pt=0)
    _add_paragraph(doc, f"4) Aderência a Política de Investimentos - Os recursos investidos, abrangendo a carteira consolidada, que representa o patrimônio total do RPPS sob gestão, estão aderentes à Política de Investimentos de {meeting['ano']}, respeitando o estabelecido na legislação em vigor e dentro dos percentuais definidos.  Considerando que as taxas ainda são negociadas acima da meta atuarial, seguimos com a estratégia de alcançar o alvo definido de 60% de alocação em Títulos Públicos.", False, space_after_pt=SECTION_GAP)


    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(10)  # controla a altura do espaço

    # Item 04
    _add_title(doc, "Item 04 – Assuntos Gerais", space_after_pt=TITLE_GAP)
    _add_paragraph(doc, STATE["assuntos"], False, WD_ALIGN_PARAGRAPH.JUSTIFY, space_after_pt=SECTION_GAP)

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(10)  # controla a altura do espaço

  # Fecho
    lav = (meeting.get("lavrador") or "").strip() or "___________________________________"
    _add_paragraph(
      doc,
      f"Nada mais havendo a tratar, foi encerrada a reunião e eu, {lav}, lavrei a presente Ata, assinada pelos membros presentes do Comitê de Investimentos.",
      False,
      space_after_pt=SECTION_GAP
)

# Espaço extra antes das assinaturas (1 linha vazia)
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(10)  # controla a altura do espaço

# Assinaturas
    _assinaturas_table(doc)


    # Saída
    buf = io.BytesIO(); doc.save(buf); buf.seek(0)
    filename = f"Ata_{meeting['numero']:03d}-{meeting['ano']}.docx"
    return StreamingResponse(buf, media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                             headers={"Content-Disposition": f"attachment; filename={filename}"})
